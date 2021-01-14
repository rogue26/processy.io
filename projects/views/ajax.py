import io

from django.http import JsonResponse
from django.template.loader import render_to_string
from django.http import HttpResponse
from django.core.files.base import ContentFile
from django.db.models import Q

from docx import Document

from projects.models import Workstream, Deliverable, Task, TeamMember
from projects.models import ScopeOfWork, Project


def project_scope_details():
    pass


def project_approach_details():
    pass


def project_deliverables_details():
    pass


def project_team_details():
    pass


def project_investment_details():
    pass


def ajax_test(request):
    if request.method == 'GET':
        project_id = request.GET['project_id']

        project = Project.objects.get(id=project_id)

        replace_dict = \
            {
                # "organization_name": project.organization.name,
                # "organization_address": project.organization.address,
                # "client_name": project.client.name,
                # "client_description": project.client.description,
                # "client_address": project.client.address,
                # "client_nickname": project.client.nickname(),
                # "project_name": project.name,
                # "today_date": datetime.date.today(),
                # "project_objectives": None,  # loop through each workstream objective as bullet point
                # "regional_scope": None,
                # "product_scope": None,
                # "business_unit_scope": None,
                # "channel_scope": None,

            }

        document = Document()
        document.add_heading('Document Title', 0)
        document.add_paragraph('A plain paragraph having some ')

        # Create in-memory buffer
        file_stream = io.BytesIO()
        # Save the .docx to the buffer
        document.save(file_stream)
        # Reset the buffer's file-pointer to the beginning of the file
        file_stream.seek(0)

        document.save(file_stream)

        # word_file = ContentFile(file_stream.getvalue().encode('utf-8'))
        word_file = ContentFile(file_stream.getvalue())
        # source_stream.close()

        # target_stream = StringIO()
        # document.save(target_stream)

        # row = ["Project ID", project_id]
        # buffer = StringIO()
        # csv_writer = csv.writer(buffer)
        # csv_writer.writerow(row)
        # csv_file = ContentFile(buffer.getvalue().encode('utf-8'))
        sow = ScopeOfWork()
        # sow.file.save('output.csv', csv_file)
        sow.file.save('output.docx', word_file)
        sow.save()
        return HttpResponse(sow.file.url)
    else:
        return HttpResponse("unsuccesful")


def update_projects_table(request):
    data = dict()
    if request.method == 'GET':
        projects = Project.objects.filter(Q(team_member__user=request.user))
        data['data'] = render_to_string(
            'projects/tables/projects_table.html',
            {'projects': projects},
            request=request
        )
        return JsonResponse(data)


def update_workstreams_table(request, project_id):
    data = dict()
    if request.method == 'GET':
        workstreams = Workstream.objects.filter(project_id=project_id)
        data['data'] = render_to_string(
            'projects/tables/workstreams_table.html',
            {'workstreams': workstreams},
            request=request
        )
        return JsonResponse(data)


def update_deliverables_table(request, project_id):
    data = dict()
    if request.method == 'GET':
        deliverables = Deliverable.objects.filter(project_id=project_id)
        data['data'] = render_to_string(
            'projects/tables/deliverables_table.html',
            {'deliverables': deliverables},
            request=request
        )
        return JsonResponse(data)


def update_tasks_table(request, project_id):
    data = dict()
    if request.method == 'GET':
        tasks = Task.objects.filter(project_id=project_id)
        data['data'] = render_to_string(
            'projects/tables/tasks_table.html',
            {'tasks': tasks},
            request=request
        )
        return JsonResponse(data)


def update_team_members_table(request, project_id):
    data = dict()
    if request.method == 'GET':
        team_members = TeamMember.objects.filter(project_id=project_id)
        data['data'] = render_to_string(
            'projects/tables/team_members_table.html',
            {'team_members': team_members},
            request=request
        )
        return JsonResponse(data)
